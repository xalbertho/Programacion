from docx import Document

# Create a new Document
doc = Document()
doc.add_heading('Comparación de Procesos Administrativos', level=1)

# Add the table
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'

# Add header row
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Aspecto'
hdr_cells[1].text = "McDonald's"
hdr_cells[2].text = 'LEGO'
hdr_cells[3].text = 'Cerveza Toro'

# Add other rows
data = [
    ("Planeación", 
     "Estrategias a largo plazo para la expansión global y sostenibilidad.",
     "Innovación de productos y colaboraciones estratégicas con franquicias. Planes de sostenibilidad.",
     "Planificación de productos que combinan cerveza y tradiciones vinícolas. Expansión en mercados internacionales."),
    ("Organización",
     "Estructura organizativa global con procesos estandarizados.",
     "División por departamentos con roles claros. Colaboración interna para innovación.",
     "Coordinación entre producción y colaboradores para el uso de barricas de Jerez."),
    ("Dirección",
     "Liderazgo en niveles múltiples, desde gerentes de tienda hasta ejecutivos globales.",
     "Cultura corporativa que fomenta creatividad y aprendizaje.",
     "Liderazgo que promueve la calidad y el respeto por la tradición y la innovación."),
    ("Control",
     "Auditorías internas, controles de calidad y cumplimiento de regulaciones.",
     "KPIs para medir productividad y satisfacción del cliente. Revisión de estándares de calidad y cumplimiento.",
     "Controles de calidad y autenticidad en el producto final, monitoreo del proceso de envejecimiento.")
]

for aspect, mcd, lego, toro in data:
    row_cells = table.add_row().cells
    row_cells[0].text = aspect
    row_cells[1].text = mcd
    row_cells[2].text = lego
    row_cells[3].text = toro

# Save the document
file_path = "/mnt/data/Comparacion_Procesos_Administrativos.docx"
doc.save(file_path)
file_path
